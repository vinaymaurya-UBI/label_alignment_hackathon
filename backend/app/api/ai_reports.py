"""
AI report endpoints: SSE streaming generation + DOCX download.
"""
import io
import json
import re
from typing import AsyncGenerator

from fastapi import APIRouter, Depends
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.core.database import get_db
from app.models import ActivityLog, Drug
from app.services.ai_service import get_ai_generator

router = APIRouter()


# ---------------------------------------------------------------------------
# SSE streaming report generation
# ---------------------------------------------------------------------------

async def _stream(db: AsyncSession, drug_id: str) -> AsyncGenerator[bytes, None]:
    def evt(data: dict) -> bytes:
        return b"data: " + json.dumps(data).encode() + b"\n\n"

    yield evt({"status": "starting", "progress": 0, "message": "Initializing..."})

    ai = get_ai_generator()
    if ai is None:
        yield evt({"status": "error", "message": "AI service not configured (missing ANTHROPIC_API_KEY)"})
        return

    try:
        yield evt({"status": "generating", "progress": 30, "message": "Generating AI-powered analysis..."})
        report_md = await ai.generate_report(db, drug_id)
        
        # Log the activity
        drug_res = await db.execute(select(Drug).where(Drug.id == drug_id))
        drug = drug_res.scalar_one_or_none()
        drug_name = drug.generic_name if drug else drug_id
        
        new_log = ActivityLog(
            type="report",
            title=f"AI Comparison Generated: {drug_name}",
            subtitle=f"Cross-country alignment report for {drug_name}",
            status="Completed",
            meta={"drug_id": drug_id}
        )
        db.add(new_log)
        await db.commit()

        yield evt({"status": "complete", "progress": 100, "message": "Done!", "report": report_md})
    except ValueError as exc:
        yield evt({"status": "error", "message": str(exc)})
    except Exception as exc:
        yield evt({"status": "error", "message": f"AI generation failed: {exc}"})


@router.post("/generate-report/{drug_id}")
async def generate_report_stream(drug_id: str, db: AsyncSession = Depends(get_db)):
    return StreamingResponse(
        _stream(db, drug_id),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ---------------------------------------------------------------------------
# DOCX download
# ---------------------------------------------------------------------------

class DocxRequest(BaseModel):
    report: str


def _add_formatted_text(paragraph, text):
    """
    Helper to add text with **bold** and *italic* support to a paragraph.
    """
    # Pattern to find **bold**, *italic*, or plain text
    # We use a non-greedy match to handle multiple occurrences
    parts = re.split(r"(\*\*.+?\*\*|\*.+? \*)", text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _markdown_to_docx(markdown: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Style the default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Handle Tables
        if stripped.startswith("|") and i + 1 < len(lines) and "|" in lines[i+1] and "---" in lines[i+1]:
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            
            if len(table_lines) >= 3: # Header, separator, at least one data row
                headers = [c.strip() for c in table_lines[0].strip("|").split("|") if c.strip()]
                data_rows = []
                for dl in table_lines[2:]: # Skip header and separator
                    row = [c.strip() for c in dl.strip("|").split("|") if c.strip()]
                    if row:
                        data_rows.append(row)
                
                if headers:
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for idx, h in enumerate(headers):
                        hdr_cells[idx].text = h
                        # Bold the header text
                        for paragraph in hdr_cells[idx].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    
                    for row_data in data_rows:
                        row_cells = table.add_row().cells
                        # Ensure we don't exceed column count
                        for idx, cell_text in enumerate(row_data[:len(headers)]):
                            # We can use our formatted text helper here too
                            p = row_cells[idx].paragraphs[0]
                            _add_formatted_text(p, cell_text)
            continue

        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(para, stripped[2:])
        elif stripped.startswith("---"):
            doc.add_paragraph("─" * 60)
        elif stripped == "":
            doc.add_paragraph("")
        else:
            para = doc.add_paragraph()
            _add_formatted_text(para, stripped)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.post("/download-docx/{drug_id}")
async def download_docx(drug_id: str, body: DocxRequest, db: AsyncSession = Depends(get_db)):
    docx_bytes = _markdown_to_docx(body.report)
    filename = f"drug_report_{drug_id}.docx"
    
    # Log the export activity
    drug_res = await db.execute(select(Drug).where(Drug.id == drug_id))
    drug = drug_res.scalar_one_or_none()
    drug_name = drug.generic_name if drug else drug_id
    
    new_log = ActivityLog(
        type="export",
        title=f"Report Exported: {drug_name}",
        subtitle=f"Downloaded regulatory comparison as DOCX",
        status="Downloaded",
        meta={"drug_id": drug_id, "format": "DOCX"}
    )
    db.add(new_log)
    await db.commit()

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# AI status
# ---------------------------------------------------------------------------

@router.get("/status")
async def ai_status():
    ai = get_ai_generator()
    if ai is None:
        return {"status": "error", "available": False, "message": "No AI provider configured"}
    
    providers = []
    if ai.gemini_enabled:
        providers.append(f"Google Gemini ({settings.GOOGLE_MODEL})")
    if ai.anthropic_client:
        providers.append(f"Anthropic Claude ({settings.ANTHROPIC_MODEL})")
    
    return {
        "status": "available", 
        "available": True, 
        "provider": " / ".join(providers) if providers else "Fallback Only"
    }
